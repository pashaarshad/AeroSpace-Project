import os
import spacy
import json
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime
import requests
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and analyze documents for knowledge extraction"""
    
    def __init__(self):
        self.setup_nlp()
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.html']
        self.processed_documents = []
        self.allowed_extensions = {"pdf", "docx", "txt"}
    
    def setup_nlp(self):
        """Initialize NLP models"""
        try:
            # Try to load spaCy model
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except OSError:
            logger.warning("spaCy model not found, using basic processing")
            self.nlp = None
    
    def process_uploaded_file(self, file) -> Dict[str, Any]:
        """Process an uploaded file and extract knowledge"""
        try:
            filename = secure_filename(file.filename)
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text from file
            text_content = self.extract_text_from_file(file, file_extension)
            
            # Process the text
            processed_data = self.process_text(text_content, filename)
            
            # Store processed document info
            doc_info = {
                'filename': filename,
                'processed_at': datetime.now().isoformat(),
                'entities_count': len(processed_data.get('entities', [])),
                'relations_count': len(processed_data.get('relations', [])),
                'word_count': len(text_content.split()),
                'file_type': file_extension
            }
            
            self.processed_documents.append(doc_info)
            
            return {
                **processed_data,
                'document_info': doc_info
            }
            
        except Exception as e:
            logger.error(f"File processing error: {str(e)}")
            raise
    
    def extract_text_from_file(self, file, file_extension: str) -> str:
        """Extract text content from different file formats"""
        try:
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(file)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(file)
            elif file_extension == '.html':
                return self.extract_text_from_html(file)
            elif file_extension in ['.txt', '.md']:
                return file.read().decode('utf-8')
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise
    
    def extract_text_from_pdf(self, file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise
    
    def extract_text_from_html(self, file) -> str:
        """Extract text from HTML file"""
        try:
            content = file.read().decode('utf-8')
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"HTML extraction error: {str(e)}")
            raise
    
    def process_text(self, text: str, source: str = "unknown") -> Dict[str, Any]:
        """Process text and extract entities and relationships"""
        try:
            # Basic text processing
            processed_data = {
                'entities': [],
                'relations': [],
                'keywords': [],
                'summary': '',
                'source': source,
                'processed_at': datetime.now().isoformat()
            }
            
            if self.nlp:
                # Advanced NLP processing
                doc = self.nlp(text)
                
                # Extract entities
                entities = self.extract_entities(doc)
                processed_data['entities'] = entities
                
                # Extract relationships
                relations = self.extract_relationships(doc)
                processed_data['relations'] = relations
                
                # Extract keywords
                keywords = self.extract_keywords(doc)
                processed_data['keywords'] = keywords
                
                # Generate summary
                summary = self.generate_summary(text)
                processed_data['summary'] = summary
                
            else:
                # Basic processing without spaCy
                processed_data.update(self.basic_text_processing(text))
            
            return processed_data
            
        except Exception as e:
            logger.error(f"Text processing error: {str(e)}")
            raise
    
    def extract_entities(self, doc) -> List[Dict[str, Any]]:
        """Extract named entities from spaCy doc"""
        entities = []
        
        for ent in doc.ents:
            entity = {
                'name': ent.text,
                'type': ent.label_,
                'description': spacy.explain(ent.label_),
                'start_char': ent.start_char,
                'end_char': ent.end_char,
                'confidence': 0.8  # Default confidence
            }
            entities.append(entity)
        
        # Remove duplicates
        unique_entities = []
        seen = set()
        
        for entity in entities:
            key = (entity['name'].lower(), entity['type'])
            if key not in seen:
                seen.add(key)
                unique_entities.append(entity)
        
        return unique_entities
    
    def extract_relationships(self, doc) -> List[Dict[str, Any]]:
        """Extract relationships between entities"""
        relationships = []
        
        # Simple relationship extraction based on dependency parsing
        for sent in doc.sents:
            for token in sent:
                if token.dep_ in ["nsubj", "dobj", "pobj"]:
                    # Find related entities
                    head = token.head
                    
                    # Check if both token and head are entities
                    token_ent = self.find_entity_for_token(token, doc.ents)
                    head_ent = self.find_entity_for_token(head, doc.ents)
                    
                    if token_ent and head_ent and token_ent != head_ent:
                        relationship = {
                            'source': token_ent.text,
                            'target': head_ent.text,
                            'type': token.dep_,
                            'description': f"{token_ent.text} {token.dep_} {head_ent.text}",
                            'confidence': 0.7
                        }
                        relationships.append(relationship)
        
        return relationships
    
    def find_entity_for_token(self, token, entities):
        """Find entity that contains the given token"""
        for ent in entities:
            if ent.start <= token.i < ent.end:
                return ent
        return None
    
    def extract_keywords(self, doc) -> List[Dict[str, Any]]:
        """Extract important keywords from text"""
        keywords = []
        
        # Extract noun phrases
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) <= 3:  # Limit to 3 words
                keyword = {
                    'text': chunk.text,
                    'type': 'noun_phrase',
                    'importance': self.calculate_keyword_importance(chunk.text, doc.text)
                }
                keywords.append(keyword)
        
        # Extract important single words
        for token in doc:
            if (token.pos_ in ["NOUN", "PROPN", "ADJ"] and 
                not token.is_stop and 
                not token.is_punct and 
                len(token.text) > 2):
                
                keyword = {
                    'text': token.text,
                    'type': token.pos_,
                    'importance': self.calculate_keyword_importance(token.text, doc.text)
                }
                keywords.append(keyword)
        
        # Sort by importance and remove duplicates
        unique_keywords = []
        seen = set()
        
        for keyword in sorted(keywords, key=lambda x: x['importance'], reverse=True):
            if keyword['text'].lower() not in seen:
                seen.add(keyword['text'].lower())
                unique_keywords.append(keyword)
        
        return unique_keywords[:50]  # Return top 50 keywords
    
    def calculate_keyword_importance(self, keyword: str, text: str) -> float:
        """Calculate importance score for a keyword"""
        # Simple TF-IDF-like scoring
        keyword_lower = keyword.lower()
        text_lower = text.lower()
        
        # Term frequency
        tf = text_lower.count(keyword_lower) / len(text_lower.split())
        
        # Length bonus (longer keywords are often more important)
        length_bonus = min(len(keyword.split()), 3) * 0.1
        
        # Position bonus (keywords at the beginning are often more important)
        position_bonus = 0.1 if text_lower.index(keyword_lower) < len(text_lower) * 0.2 else 0
        
        return tf + length_bonus + position_bonus
    
    def generate_summary(self, text: str, max_sentences: int = 3) -> str:
        """Generate a simple extractive summary"""
        try:
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
            
            if len(sentences) <= max_sentences:
                return '. '.join(sentences)
            
            # Simple scoring based on sentence length and position
            scored_sentences = []
            
            for i, sentence in enumerate(sentences):
                # Position score (earlier sentences are more important)
                position_score = 1.0 / (i + 1)
                
                # Length score (medium-length sentences are preferred)
                length_score = min(len(sentence.split()) / 20, 1.0)
                
                total_score = position_score * 0.6 + length_score * 0.4
                
                scored_sentences.append((sentence, total_score))
            
            # Select top sentences
            top_sentences = sorted(scored_sentences, key=lambda x: x[1], reverse=True)[:max_sentences]
            
            # Maintain original order
            summary_sentences = []
            for sentence, _ in top_sentences:
                original_index = sentences.index(sentence)
                summary_sentences.append((original_index, sentence))
            
            summary_sentences.sort(key=lambda x: x[0])
            
            return '. '.join([sentence for _, sentence in summary_sentences])
            
        except Exception as e:
            logger.error(f"Summary generation error: {str(e)}")
            return text[:500] + "..." if len(text) > 500 else text
    
    def basic_text_processing(self, text: str) -> Dict[str, Any]:
        """Basic text processing without advanced NLP"""
        try:
            words = text.split()
            
            # Extract potential entities (capitalized words)
            entities = []
            for word in words:
                if word[0].isupper() and len(word) > 2:
                    entities.append({
                        'name': word,
                        'type': 'UNKNOWN',
                        'description': 'Potential entity (capitalized word)',
                        'confidence': 0.5
                    })
            
            # Extract keywords (frequent words)
            word_freq = {}
            for word in words:
                word_lower = word.lower()
                if len(word_lower) > 3 and word_lower.isalpha():
                    word_freq[word_lower] = word_freq.get(word_lower, 0) + 1
            
            keywords = []
            for word, freq in sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]:
                keywords.append({
                    'text': word,
                    'type': 'WORD',
                    'importance': freq / len(words)
                })
            
            return {
                'entities': entities[:50],  # Limit to 50 entities
                'relations': [],  # No relations without advanced NLP
                'keywords': keywords,
                'summary': self.generate_summary(text)
            }
            
        except Exception as e:
            logger.error(f"Basic text processing error: {str(e)}")
            return {
                'entities': [],
                'relations': [],
                'keywords': [],
                'summary': ''
            }
    
    def search_documents(self, query: str) -> List[Dict[str, Any]]:
        """Search processed documents for relevant information"""
        try:
            results = []
            query_lower = query.lower()
            
            # This is a placeholder - in a real implementation, you'd search
            # through stored document content, possibly using vector search
            
            # For now, return some mock results
            mock_results = [
                {
                    'title': 'AI Research Document',
                    'snippet': 'This document discusses artificial intelligence and machine learning approaches...',
                    'relevance': 0.9,
                    'source': 'research_paper.pdf'
                },
                {
                    'title': 'Knowledge Graph Guide',
                    'snippet': 'A comprehensive guide to building and using knowledge graphs...',
                    'relevance': 0.8,
                    'source': 'kg_guide.docx'
                }
            ]
            
            # Filter results based on query
            for result in mock_results:
                if any(word in result['snippet'].lower() for word in query_lower.split()):
                    results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Document search error: {str(e)}")
            return []
    
    def get_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        try:
            total_docs = len(self.processed_documents)
            
            if total_docs == 0:
                return {
                    'total_documents': 0,
                    'total_entities': 0,
                    'total_relations': 0,
                    'file_types': {}
                }
            
            total_entities = sum(doc['entities_count'] for doc in self.processed_documents)
            total_relations = sum(doc['relations_count'] for doc in self.processed_documents)
            
            # Count file types
            file_types = {}
            for doc in self.processed_documents:
                file_type = doc['file_type']
                file_types[file_type] = file_types.get(file_type, 0) + 1
            
            return {
                'total_documents': total_docs,
                'total_entities': total_entities,
                'total_relations': total_relations,
                'file_types': file_types,
                'processed_documents': self.processed_documents
            }
            
        except Exception as e:
            logger.error(f"Stats error: {str(e)}")
            return {}
    
    def allowed_file(self, filename: str) -> bool:
        """Check if the file extension is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in self.allowed_extensions

    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract knowledge from the document"""
        try:
            if file_path.endswith(".pdf"):
                return self.process_pdf(file_path)
            elif file_path.endswith(".docx"):
                return self.process_docx(file_path)
            elif file_path.endswith(".txt"):
                return self.process_txt(file_path)
            else:
                logger.error("Unsupported file type")
                return []
        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            return []

    def process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from PDF"""
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join(page.extract_text() for page in reader.pages)
                return self.extract_entities(text)
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            return []

    def process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return self.extract_entities(text)
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            return []

    def process_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from TXT"""
        try:
            with open(file_path, "r") as f:
                text = f.read()
                return self.extract_entities(text)
        except Exception as e:
            logger.error(f"Failed to process TXT: {e}")
            return []

    def extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract entities using NLP"""
        try:
            doc = self.nlp(text)
            entities = [
                {
                    "text": ent.text,
                    "label": ent.label_
                }
                for ent in doc.ents
            ]
            return entities
        except Exception as e:
            logger.error(f"Failed to extract entities: {e}")
            return []

# Example usage and testing
if __name__ == '__main__':
    processor = DocumentProcessor()
    
    # Test text processing
    sample_text = """
    Artificial Intelligence (AI) is a branch of computer science that aims to create 
    intelligent machines. Machine Learning is a subset of AI that enables computers 
    to learn without being explicitly programmed. Deep Learning uses neural networks 
    to model and understand complex patterns.
    """
    
    result = processor.process_text(sample_text, "sample_text")
    print(f"Processed text result: {json.dumps(result, indent=2)}")
    
    # Test document search
    search_results = processor.search_documents("artificial intelligence")
    print(f"Search results: {json.dumps(search_results, indent=2)}")
    
    # Test stats
    stats = processor.get_stats()
    print(f"Stats: {json.dumps(stats, indent=2)}")
